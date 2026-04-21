#!/usr/bin/env python3
"""
Release Readiness Dashboard - Export Server
Handles Word and PDF export generation
"""

import json
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors


class ReleaseReadinessExporter:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._add_custom_styles()

    def _add_custom_styles(self):
        """Add custom styles for PDF"""
        self.styles.add(
            ParagraphStyle(
                name="CustomTitle",
                parent=self.styles["Heading1"],
                fontSize=24,
                textColor=colors.HexColor("#1F2937"),
                spaceAfter=12,
                fontName="Helvetica-Bold",
            )
        )
        self.styles.add(
            ParagraphStyle(
                name="CustomHeading",
                parent=self.styles["Heading2"],
                fontSize=14,
                textColor=colors.HexColor("#374151"),
                spaceAfter=12,
                spaceBefore=6,
            )
        )

    def export_to_word(self, data: dict, output_path: str = "release_readiness.docx"):
        """Generate Word document from release readiness data"""
        doc = Document()

        # Title
        title = doc.add_heading("🚀 Release Readiness Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Release Info
        info_section = doc.add_heading("Release Information", level=1)
        doc.add_paragraph(f"Release Name: {data['releaseName']}")
        doc.add_paragraph(f"Target Date: {data['releaseDate']}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Summary Section
        doc.add_heading("Executive Summary", level=1)
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = "Light Grid Accent 1"

        summary_data = data["summary"]
        summary_table.cell(0, 0).text = "Overall Progress"
        summary_table.cell(0, 1).text = f"{summary_data['overallProgress']}%"

        summary_table.cell(1, 0).text = "Team Sign-offs"
        summary_table.cell(1, 1).text = f"{summary_data['signOffsApproved']}/{summary_data['signOffsTotal']}"

        summary_table.cell(2, 0).text = "Documentation"
        summary_table.cell(2, 1).text = f"{summary_data['docsCompleted']}/{summary_data['docsTotal']}"

        summary_table.cell(3, 0).text = "Risk Level"
        summary_table.cell(3, 1).text = summary_data["riskLevel"]

        summary_table.cell(4, 0).text = "Status"
        status = (
            "🟢 READY FOR LAUNCH"
            if summary_data["overallProgress"] >= 80
            else "🟡 IN PROGRESS"
        )
        summary_table.cell(4, 1).text = status

        # Team Sign-offs Section
        doc.add_heading("Team Sign-offs", level=1)
        signoffs_table = doc.add_table(rows=1, cols=4)
        signoffs_table.style = "Light Grid Accent 1"
        hdr_cells = signoffs_table.rows[0].cells
        hdr_cells[0].text = "Team"
        hdr_cells[1].text = "Owner"
        hdr_cells[2].text = "Status"
        hdr_cells[3].text = "Comment"

        for signoff in data["teamSignOffs"]:
            row_cells = signoffs_table.add_row().cells
            row_cells[0].text = signoff["team"]
            row_cells[1].text = signoff["owner"]
            row_cells[2].text = signoff["status"].upper()
            row_cells[3].text = signoff["comment"] or "—"

        # Risk Assessment Section
        doc.add_heading("Risk Assessment", level=1)
        risk_table = doc.add_table(rows=1, cols=4)
        risk_table.style = "Light Grid Accent 1"
        hdr_cells = risk_table.rows[0].cells
        hdr_cells[0].text = "Risk"
        hdr_cells[1].text = "Severity"
        hdr_cells[2].text = "Probability"
        hdr_cells[3].text = "Mitigation"

        for risk in data["risks"]:
            row_cells = risk_table.add_row().cells
            row_cells[0].text = risk["risk"]
            row_cells[1].text = risk["severity"].upper()
            row_cells[2].text = risk["probability"].upper()
            row_cells[3].text = risk["mitigation"]

        # Documentation Checklist
        doc.add_heading("Documentation Checklist", level=1)
        for doc_item in data["documentation"]:
            status = "✓" if doc_item["completed"] else "✗"
            doc.add_paragraph(
                f"{status} {doc_item['item']}", style="List Bullet"
            )

        # Save document
        doc.save(output_path)
        return output_path

    def export_to_pdf(self, data: dict, output_path: str = "release_readiness.pdf"):
        """Generate PDF document from release readiness data"""
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []

        # Title
        story.append(Paragraph("🚀 Release Readiness Report", self.styles["CustomTitle"]))
        story.append(Spacer(1, 0.2 * inch))

        # Release Info
        story.append(Paragraph("Release Information", self.styles["CustomHeading"]))
        release_info = [
            ["Release Name:", data["releaseName"]],
            ["Target Date:", data["releaseDate"]],
            ["Generated:", datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
        ]
        release_table = Table(release_info, colWidths=[1.5 * inch, 4.5 * inch])
        release_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F3F4F6")),
                    ("TEXTCOLOR", (0, 0), (-1, -1), colors.black),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ]
            )
        )
        story.append(release_table)
        story.append(Spacer(1, 0.2 * inch))

        # Summary
        story.append(Paragraph("Executive Summary", self.styles["CustomHeading"]))
        summary_data = data["summary"]
        summary_info = [
            ["Overall Progress", f"{summary_data['overallProgress']}%"],
            ["Team Sign-offs", f"{summary_data['signOffsApproved']}/{summary_data['signOffsTotal']}"],
            ["Documentation", f"{summary_data['docsCompleted']}/{summary_data['docsTotal']}"],
            ["Risk Level", summary_data["riskLevel"]],
            [
                "Status",
                "🟢 READY FOR LAUNCH"
                if summary_data["overallProgress"] >= 80
                else "🟡 IN PROGRESS",
            ],
        ]
        summary_table = Table(summary_info, colWidths=[2 * inch, 4 * inch])
        summary_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#3B82F6")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F9FAFB")]),
                ]
            )
        )
        story.append(summary_table)
        story.append(Spacer(1, 0.3 * inch))

        # Team Sign-offs
        story.append(Paragraph("Team Sign-offs", self.styles["CustomHeading"]))
        signoff_data = [["Team", "Owner", "Status", "Comment"]]
        for signoff in data["teamSignOffs"]:
            signoff_data.append(
                [
                    signoff["team"],
                    signoff["owner"],
                    signoff["status"].upper(),
                    signoff["comment"] or "—",
                ]
            )
        signoff_table = Table(signoff_data, colWidths=[1.5 * inch, 1.5 * inch, 1 * inch, 2 * inch])
        signoff_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#10B981")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ]
            )
        )
        story.append(signoff_table)
        story.append(PageBreak())

        # Risk Assessment
        story.append(Paragraph("Risk Assessment", self.styles["CustomHeading"]))
        risk_data = [["Risk", "Severity", "Probability", "Mitigation"]]
        for risk in data["risks"]:
            risk_data.append(
                [
                    risk["risk"],
                    risk["severity"].upper(),
                    risk["probability"].upper(),
                    risk["mitigation"],
                ]
            )
        risk_table = Table(risk_data, colWidths=[1.8 * inch, 1 * inch, 1 * inch, 2.2 * inch])
        risk_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F97316")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ]
            )
        )
        story.append(risk_table)
        story.append(Spacer(1, 0.2 * inch))

        # Documentation
        story.append(Paragraph("Documentation Checklist", self.styles["CustomHeading"]))
        doc_items = []
        for doc_item in data["documentation"]:
            status = "✓" if doc_item["completed"] else "✗"
            doc_items.append(f"{status} {doc_item['item']}")

        for item in doc_items:
            story.append(Paragraph(item, self.styles["Normal"]))

        # Build PDF
        doc.build(story)
        return output_path


def main():
    # Example usage
    exporter = ReleaseReadinessExporter()

    # Sample data
    sample_data = {
        "releaseName": "Q2 Feature Release",
        "releaseDate": "2026-06-30",
        "teamSignOffs": [
            {"team": "Engineering", "owner": "John Smith", "status": "approved", "comment": "All tests passing"},
            {"team": "Product", "owner": "Sarah Chen", "status": "pending", "comment": ""},
            {"team": "QA", "owner": "Mike Johnson", "status": "approved", "comment": "UAT complete"},
        ],
        "risks": [
            {
                "risk": "API rate limiting not tested",
                "severity": "high",
                "probability": "medium",
                "mitigation": "Load testing scheduled",
            },
        ],
        "documentation": [
            {"item": "Release notes drafted", "completed": True},
            {"item": "API documentation updated", "completed": False},
        ],
        "summary": {
            "overallProgress": 70,
            "signOffsApproved": 2,
            "signOffsTotal": 3,
            "docsCompleted": 1,
            "docsTotal": 2,
            "riskLevel": "MEDIUM",
        },
    }

    # Export to Word
    print("Generating Word document...")
    word_path = exporter.export_to_word(sample_data, "release_readiness.docx")
    print(f"✓ Word document saved: {word_path}")

    # Export to PDF
    print("Generating PDF...")
    pdf_path = exporter.export_to_pdf(sample_data, "release_readiness.pdf")
    print(f"✓ PDF saved: {pdf_path}")


if __name__ == "__main__":
    main()
